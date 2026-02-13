from __future__ import annotations

import hashlib
import os
import re
from pathlib import Path
from importlib.resources import files as pkg_files

import typer
from openai import OpenAI
from rich import print
from rich.console import Console

app = typer.Typer(no_args_is_help=True)
console = Console()

# Filename ordering helper: prefers patterns like 3.7.1, 3.7.2, etc.
NUM = re.compile(r"(\d+)(?:\.(\d+))?(?:\.(\d+))?")

SUPPORTED_EXTS = {".txt", ".md", ".srt", ".vtt"}
DEFAULT_CHUNK_CHARS = 120_000


def sort_key(p: Path):
    """
    Prefer numeric ordering in filenames like 3.7.1, 3.7.2, etc.
    Fall back to modified time (older first) when no numeric pattern exists.
    """
    m = NUM.search(p.stem)
    if m:
        parts = tuple(int(x) if x else -1 for x in m.groups())
        return (0, parts, p.name.lower())
    return (1, p.stat().st_mtime, p.name.lower())


def list_transcripts(folder: Path) -> list[Path]:
    files = [p for p in folder.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS]
    files.sort(key=sort_key)
    return files


def read_text(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="ignore").strip()


def load_default_prompt() -> str:
    # Bundled file inside the installed package
    prompt_path = pkg_files("t2md").joinpath("default_prompt.md")
    return prompt_path.read_text(encoding="utf-8").strip()


def derive_module_name(folder: Path) -> str:
    return folder.name


def sanitize_name(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._-")
    return cleaned or "module"


def split_large_text(text: str, max_chars: int) -> list[str]:
    """
    Split oversized text by paragraph boundaries where possible, and hard-split as fallback.
    """
    if len(text) <= max_chars:
        return [text]

    pieces: list[str] = []
    current: list[str] = []
    current_len = 0

    for para in text.split("\n\n"):
        if not para:
            continue
        addition = para if not current else f"\n\n{para}"
        if current_len + len(addition) <= max_chars:
            current.append(addition)
            current_len += len(addition)
            continue

        if current:
            pieces.append("".join(current))
            current = []
            current_len = 0

        if len(para) <= max_chars:
            current = [para]
            current_len = len(para)
            continue

        for i in range(0, len(para), max_chars):
            chunk = para[i : i + max_chars]
            if chunk:
                pieces.append(chunk)

    if current:
        pieces.append("".join(current))

    return pieces if pieces else [text]


def build_transcript_chunks(files: list[Path], chunk_chars: int) -> list[str]:
    units: list[str] = []
    file_piece_limit = max(10_000, chunk_chars - 4_000)

    for file_path in files:
        raw = read_text(file_path)
        if not raw:
            continue
        segments = split_large_text(raw, file_piece_limit)
        total = len(segments)
        for i, segment in enumerate(segments, start=1):
            suffix = "" if total == 1 else f" (part {i}/{total})"
            units.append(f"\n\n---\n\n## SOURCE FILE: {file_path.name}{suffix}\n\n{segment}\n")

    if not units:
        return []

    chunks: list[str] = []
    current = ""
    for unit in units:
        candidate = unit if not current else current + unit
        if current and len(candidate) > chunk_chars:
            chunks.append(current.strip())
            current = unit
        else:
            current = candidate

    if current:
        chunks.append(current.strip())
    return chunks


def build_single_prompt(module_name: str, prompt_rules: str, transcripts: str) -> str:
    return f"""
You will be given:
1) PROMPT_RULES (how to transform)
2) TRANSCRIPTS (raw content)

Return EXACTLY two top-level Markdown sections:

# {module_name} — Executive Summary
- thesis (2–4 sentences)
- 5–10 key concepts
- 2–5 examples/case studies
- what to remember (3–7 bullets)

# {module_name} — Reading
Follow PROMPT_RULES to create a readable textbook-style markdown:
- include a TOC
- clean headings
- preserve important examples
- remove transcript artifacts
- end with a synthesis summary

PROMPT_RULES:
{prompt_rules}

TRANSCRIPTS:
{transcripts}
""".strip()


def build_chunk_prompt(module_name: str, prompt_rules: str, chunk_text: str, chunk_idx: int, chunk_total: int) -> str:
    return f"""
You are processing chunk {chunk_idx} of {chunk_total} for module "{module_name}".
Extract high-fidelity notes from this chunk only. Keep details that must survive into the final textbook output.

Return Markdown using this structure exactly:

# Chunk {chunk_idx} Notes
## Core Concepts
- 5-12 bullets with concise detail
## Important Examples
- examples/case studies from this chunk
## Definitions and Terms
- terms with brief definitions
## Keep-for-Final
- facts/insights that should appear in the final executive summary or reading

Rules:
- Do not invent details.
- Preserve chronology and source-specific context when relevant.
- Keep output concise but information-dense.

PROMPT_RULES:
{prompt_rules}

TRANSCRIPT_CHUNK:
{chunk_text}
""".strip()


def build_reduce_prompt(module_name: str, prompt_rules: str, chunk_notes: str) -> str:
    return f"""
You will be given:
1) PROMPT_RULES
2) CHUNK_NOTES produced from sequential transcript chunks

Synthesize them into a single coherent result.
Return EXACTLY two top-level Markdown sections:

# {module_name} — Executive Summary
- thesis (2–4 sentences)
- 5–10 key concepts
- 2–5 examples/case studies
- what to remember (3–7 bullets)

# {module_name} — Reading
Follow PROMPT_RULES to create a readable textbook-style markdown:
- include a TOC
- clean headings
- preserve important examples
- remove transcript artifacts
- end with a synthesis summary

Rules:
- Merge overlapping ideas and remove duplication.
- Preserve chronology where it improves understanding.
- Keep final structure clean and publication-ready.

PROMPT_RULES:
{prompt_rules}

CHUNK_NOTES:
{chunk_notes}
""".strip()


def compute_run_hash(
    files: list[Path], module_name: str, model: str, chunk_chars: int, prompt_rules: str, format_name: str
) -> str:
    hasher = hashlib.sha256()
    hasher.update(module_name.encode("utf-8"))
    hasher.update(model.encode("utf-8"))
    hasher.update(str(chunk_chars).encode("utf-8"))
    hasher.update(format_name.encode("utf-8"))
    hasher.update(prompt_rules.encode("utf-8"))
    for p in files:
        stat = p.stat()
        hasher.update(str(p.resolve()).encode("utf-8"))
        hasher.update(str(stat.st_size).encode("utf-8"))
        hasher.update(str(stat.st_mtime_ns).encode("utf-8"))
    return hasher.hexdigest()[:12]


def generate_with_model(client: OpenAI, model: str, prompt: str) -> str:
    resp = client.responses.create(model=model, input=prompt)
    output = (resp.output_text or "").strip()
    if not output:
        raise typer.BadParameter("Model returned an empty response. Try rerunning with a different model.")
    return output


def write_docx_from_markdown(md_text: str, docx_path: Path) -> None:
    """
    Simple Markdown -> DOCX:
      - # / ## / ### / #### become Word headings
      - "- " bullets -> List Bullet
      - "1. " numbered -> List Number
      - "> " blockquote -> Intense Quote
      - everything else -> normal paragraph

    Notes:
      - This is intentionally minimal (no tables, code fencing, inline bold parsing).
      - Produces a clean, readable DOCX that can be exported to PDF in Word/Google Docs.
    """
    from docx import Document  # type: ignore

    doc = Document()

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        # Preserve spacing
        if not line.strip():
            doc.add_paragraph("")
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            continue

        # Bullets
        if line.lstrip().startswith("- "):
            doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
            continue

        # Numbered lists
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Number")
            continue

        # Blockquotes
        if line.startswith("> "):
            doc.add_paragraph(line[2:].strip(), style="Intense Quote")
            continue

        # Default paragraph
        doc.add_paragraph(line)

    doc.save(str(docx_path))


def escape_latex(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    out = text
    for k, v in replacements.items():
        out = out.replace(k, v)
    return out


def apply_emphasis(text: str) -> str:
    """
    Convert inline **bold** and *italic* while escaping other text.
    This is intentionally minimal and non-nesting.
    """
    pattern = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")
    out: list[str] = []
    last = 0
    for m in pattern.finditer(text):
        out.append(escape_latex(text[last : m.start()]))
        if m.group(1) is not None:
            out.append(r"\textbf{" + escape_latex(m.group(1)) + r"}")
        else:
            out.append(r"\textit{" + escape_latex(m.group(2)) + r"}")
        last = m.end()
    out.append(escape_latex(text[last:]))
    return "".join(out)


def convert_inline_md(text: str) -> str:
    """
    Convert minimal inline Markdown:
      - `code` -> \\texttt{...}
      - **bold** -> \\textbf{...}
      - *italic* -> \\textit{...}

    Unmatched backticks are treated as literal text.
    """
    parts = text.split("`")
    if len(parts) == 1:
        return apply_emphasis(text)

    if len(parts) % 2 == 0:
        parts[-2] = parts[-2] + "`" + parts[-1]
        parts = parts[:-1]

    out: list[str] = []
    for i, part in enumerate(parts):
        if i % 2 == 1:
            out.append(r"\texttt{" + escape_latex(part) + r"}")
        else:
            out.append(apply_emphasis(part))
    return "".join(out)


def write_latex_from_markdown(md_text: str, tex_path: Path, title: str) -> None:
    """
    Minimal Markdown -> LaTeX:
      - # / ## / ### / #### become sectioning commands
      - "- " bullets -> itemize
      - "1. " numbered -> enumerate
      - "> " blockquote -> quote
      - fenced code blocks -> verbatim
      - everything else -> paragraph text

    Notes:
      - Intentionally minimal; focuses on compile-ready LaTeX.
      - Inline Markdown formatting is not transformed.
    """
    lines = md_text.splitlines()
    out: list[str] = []

    out.append(r"\documentclass[11pt]{article}")
    out.append(r"\usepackage[utf8]{inputenc}")
    out.append(r"\usepackage[T1]{fontenc}")
    out.append(r"\usepackage{lmodern}")
    out.append(r"\usepackage{geometry}")
    out.append(r"\usepackage{hyperref}")
    out.append(r"\usepackage{enumitem}")
    out.append(r"\usepackage{parskip}")
    out.append(r"\geometry{margin=1in}")
    out.append(r"\setlist{noitemsep}")
    out.append(r"\title{" + escape_latex(title) + r"}")
    out.append(r"\date{}")
    out.append(r"\begin{document}")
    out.append(r"\maketitle")
    out.append(r"\tableofcontents")
    out.append("")

    in_code = False
    in_quote = False
    list_mode: str | None = None  # "itemize" or "enumerate"

    def close_list() -> None:
        nonlocal list_mode
        if list_mode:
            out.append(r"\end{" + list_mode + r"}")
            list_mode = None

    def close_quote() -> None:
        nonlocal in_quote
        if in_quote:
            out.append(r"\end{quote}")
            in_quote = False

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        if line.strip().startswith("```"):
            close_list()
            close_quote()
            if not in_code:
                out.append(r"\begin{verbatim}")
                in_code = True
            else:
                out.append(r"\end{verbatim}")
                in_code = False
            continue

        if in_code:
            out.append(line)
            continue

        if not line.strip():
            close_list()
            close_quote()
            out.append("")
            continue

        if line.startswith("# "):
            close_list()
            close_quote()
            out.append(r"\section{" + convert_inline_md(line[2:].strip()) + r"}")
            continue
        if line.startswith("## "):
            close_list()
            close_quote()
            out.append(r"\subsection{" + convert_inline_md(line[3:].strip()) + r"}")
            continue
        if line.startswith("### "):
            close_list()
            close_quote()
            out.append(r"\subsubsection{" + convert_inline_md(line[4:].strip()) + r"}")
            continue
        if line.startswith("#### "):
            close_list()
            close_quote()
            out.append(r"\paragraph{" + convert_inline_md(line[5:].strip()) + r"}")
            continue

        if line.startswith("> "):
            close_list()
            if not in_quote:
                out.append(r"\begin{quote}")
                in_quote = True
            out.append(convert_inline_md(line[2:].strip()))
            continue
        if in_quote:
            close_quote()

        if line.lstrip().startswith("- "):
            item = line.lstrip()[2:].strip()
            if list_mode != "itemize":
                close_list()
                out.append(r"\begin{itemize}")
                list_mode = "itemize"
            out.append(r"\item " + convert_inline_md(item))
            continue

        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            item = m.group(1).strip()
            if list_mode != "enumerate":
                close_list()
                out.append(r"\begin{enumerate}")
                list_mode = "enumerate"
            out.append(r"\item " + convert_inline_md(item))
            continue

        close_list()
        out.append(convert_inline_md(line))

    close_list()
    close_quote()
    if in_code:
        out.append(r"\end{verbatim}")

    out.append(r"\end{document}")
    tex_path.write_text("\n".join(out), encoding="utf-8")


@app.command("run")
def run(
    folder: Path = typer.Argument(..., help="Folder containing transcript files"),
    module: str | None = typer.Option(None, "--module", help="Module name (default: folder name)"),
    out: Path = typer.Option(Path("./outputs"), "--out", help="Output directory"),
    prompt: Path | None = typer.Option(None, "--prompt", help="Optional prompt markdown file override"),
    model: str = typer.Option("gpt-4.1-mini", "--model", help="Model to use"),
    format: str = typer.Option("md", "--format", help="Output format: md, docx, or tex"),
    chunk_chars: int = typer.Option(
        DEFAULT_CHUNK_CHARS,
        "--chunk-chars",
        help="Approx character budget per chunk before auto map/reduce is used.",
    ),
    resume: bool = typer.Option(
        True,
        "--resume/--no-resume",
        help="Reuse cached chunk notes from prior runs with unchanged inputs.",
    ),
):
    """
    Generate a combined output containing:
      1) Executive Summary
      2) Textbook-style Reading

    For large inputs, transcript content is auto-chunked and processed via
    map/reduce before producing one final output.
    Output can be Markdown (.md), Word (.docx), or LaTeX (.tex).
    """
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise typer.BadParameter(
            "OPENAI_API_KEY is not set. Add it to ~/.zshrc (recommended) or export it in your shell."
        )

    folder = folder.expanduser().resolve()
    out = out.expanduser().resolve()
    out.mkdir(parents=True, exist_ok=True)

    if not folder.exists():
        raise typer.BadParameter(f"Folder not found: {folder}")

    module_name = module or derive_module_name(folder)

    # Prompt rules
    if prompt:
        prompt_path = prompt.expanduser().resolve()
        if not prompt_path.exists():
            raise typer.BadParameter(f"Prompt file not found: {prompt_path}")
        prompt_rules = read_text(prompt_path)
    else:
        prompt_rules = load_default_prompt()

    if chunk_chars < 20_000:
        raise typer.BadParameter("--chunk-chars must be >= 20000.")

    # Gather transcripts
    files = list_transcripts(folder)
    if not files:
        raise typer.BadParameter("No transcript files found (.txt/.md/.srt/.vtt).")

    chunks = build_transcript_chunks(files, chunk_chars)
    if not chunks:
        raise typer.BadParameter("All transcript files were empty after reading.")

    client = OpenAI(api_key=api_key)

    # Write output
    fmt = format.lower().strip()
    if fmt not in {"md", "docx", "tex"}:
        raise typer.BadParameter("--format must be 'md', 'docx', or 'tex'")

    if len(chunks) == 1:
        user_prompt = build_single_prompt(module_name, prompt_rules, chunks[0])
        with console.status("Generating output..."):
            output = generate_with_model(client, model, user_prompt)
    else:
        run_hash = compute_run_hash(files, module_name, model, chunk_chars, prompt_rules, fmt)
        cache_dir = out / ".t2md_runs" / f"{sanitize_name(module_name)}_{run_hash}"
        cache_dir.mkdir(parents=True, exist_ok=True)

        print(
            f"[cyan]Auto-chunking enabled:[/cyan] {len(chunks)} chunks "
            f"(~{chunk_chars} chars/chunk)."
        )
        print(f"[cyan]Chunk cache:[/cyan] {cache_dir}")

        chunk_notes: list[str] = []
        for idx, chunk_text in enumerate(chunks, start=1):
            checkpoint_file = cache_dir / f"chunk_{idx:03d}.md"
            if resume and checkpoint_file.exists():
                note = read_text(checkpoint_file)
                if note:
                    print(f"[yellow]Reused checkpoint[/yellow] chunk {idx}/{len(chunks)}")
                    chunk_notes.append(f"\n\n---\n\n# CHUNK {idx} NOTES\n\n{note}\n")
                    continue

            chunk_prompt = build_chunk_prompt(module_name, prompt_rules, chunk_text, idx, len(chunks))
            with console.status(f"Generating chunk notes {idx}/{len(chunks)}..."):
                note = generate_with_model(client, model, chunk_prompt)
            checkpoint_file.write_text(note, encoding="utf-8")
            chunk_notes.append(f"\n\n---\n\n# CHUNK {idx} NOTES\n\n{note}\n")

        reduce_prompt = build_reduce_prompt(module_name, prompt_rules, "".join(chunk_notes).strip())
        with console.status("Merging chunk notes into final output..."):
            output = generate_with_model(client, model, reduce_prompt)

    if fmt == "md":
        out_file = out / f"{module_name}_All.md"
        out_file.write_text(output, encoding="utf-8")
    elif fmt == "docx":
        out_file = out / f"{module_name}_All.docx"
        write_docx_from_markdown(output, out_file)
    else:
        out_file = out / f"{module_name}_All.tex"
        write_latex_from_markdown(output, out_file, module_name)

    print("[cyan]Processed files (in order):[/cyan]")
    for f in files:
        print(f"  - {f.name}")
    print(f"\n[green]Wrote:[/green] {out_file}")


@app.command("doctor")
def doctor():
    """Quick sanity checks."""
    ok = True
    if not os.getenv("OPENAI_API_KEY"):
        ok = False
        print("[red]Missing:[/red] OPENAI_API_KEY environment variable")
    else:
        print("[green]OK:[/green] OPENAI_API_KEY is set")

    if ok:
        print("[green]All good.[/green]")
    else:
        raise typer.Exit(code=1)
