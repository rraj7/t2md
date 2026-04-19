"""Tests for PDF and DOCX input readers."""
import textwrap
from pathlib import Path

import pytest

from t2md.cli import _read_docx, _read_pdf, read_text


def test_read_text_plain(tmp_path):
    f = tmp_path / "notes.txt"
    f.write_text("  hello world  ", encoding="utf-8")
    assert read_text(f) == "hello world"


def test_read_pdf_extracts_text(tmp_path):
    pytest.importorskip("pdfplumber")
    reportlab = pytest.importorskip("reportlab")
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_path = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(72, 720, "Deep learning lecture notes.")
    c.save()

    text = _read_pdf(pdf_path)
    assert "Deep learning" in text


def test_read_pdf_missing_pdfplumber(tmp_path, monkeypatch):
    import builtins
    real_import = builtins.__import__

    def no_pdfplumber(name, *args, **kwargs):
        if name == "pdfplumber":
            raise ImportError("mocked missing")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", no_pdfplumber)
    dummy = tmp_path / "dummy.pdf"
    dummy.write_bytes(b"%PDF-1.4")

    with pytest.raises(RuntimeError, match="pdfplumber is required"):
        _read_pdf(dummy)


def test_read_docx_extracts_paragraphs(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Introduction to transformers.")
    doc.add_paragraph("")
    doc.add_paragraph("Attention is all you need.")
    path = tmp_path / "notes.docx"
    doc.save(str(path))

    text = _read_docx(path)
    assert "Introduction to transformers." in text
    assert "Attention is all you need." in text


def test_read_text_routes_to_pdf(tmp_path, monkeypatch):
    called = []

    def fake_read_pdf(p):
        called.append(p)
        return "pdf content"

    monkeypatch.setattr("t2md.cli._read_pdf", fake_read_pdf)
    f = tmp_path / "lecture.pdf"
    f.write_bytes(b"%PDF-1.4")
    result = read_text(f)
    assert result == "pdf content"
    assert called[0] == f


def test_read_text_routes_to_docx(tmp_path, monkeypatch):
    called = []

    def fake_read_docx(p):
        called.append(p)
        return "docx content"

    monkeypatch.setattr("t2md.cli._read_docx", fake_read_docx)
    from docx import Document
    doc = Document()
    doc.add_paragraph("test")
    f = tmp_path / "notes.docx"
    doc.save(str(f))
    result = read_text(f)
    assert result == "docx content"
    assert called[0] == f
